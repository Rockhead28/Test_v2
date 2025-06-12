import io
from typing import Dict, Any, Optional, List
from copy import deepcopy
import streamlit as st
from docx import Document
from docx.table import _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# --- HELPER FUNCTIONS ---
# These helpers are well-written and do not need changes.
def replace_text_in_paragraph(paragraph: Paragraph, key: str, value: Any):
    """Replaces all occurrences of a key in a paragraph's runs."""
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, str(value))

def copy_paragraph_formatting(source_paragraph: Paragraph, target_paragraph: Paragraph):
    target_paragraph.style = source_paragraph.style
    p_format = source_paragraph.paragraph_format
    target_p_format = target_paragraph.paragraph_format
    target_p_format.alignment = p_format.alignment
    target_p_format.space_after = p_format.space_after
    target_p_format.space_before = p_format.space_before
    target_p_format.left_indent = p_format.left_indent
    target_p_format.right_indent = p_format.right_indent

def replace_with_bullet_points(paragraph: Paragraph, key: str, bullet_points: List[str]):
    """Replaces a placeholder with a list of bullet points, preserving formatting."""
    if key not in paragraph.text:
        return
        
    if not bullet_points:
        replace_text_in_paragraph(paragraph, key, "")
        return

    # Find a run with the placeholder to use as a style template
    template_run = next((run for run in paragraph.runs if key in run.text), None)
    if not template_run:
        return

    # Replace the placeholder key with the first bullet point in the original paragraph
    replace_text_in_paragraph(paragraph, key, "• " + bullet_points[0])

    # Add subsequent bullet points as new paragraphs right after the current one
    current_p_element = paragraph._p
    for point in bullet_points[1:]:
        new_p = paragraph._parent.add_paragraph()
        copy_paragraph_formatting(paragraph, new_p)
        new_run = new_p.add_run("• " + point)
        if template_run:
            new_run.font.bold = template_run.font.bold
            new_run.font.italic = template_run.font.italic
            new_run.font.name = template_run.font.name
            new_run.font.size = template_run.font.size
        current_p_element.addnext(new_p._p)
        current_p_element = new_p._p

# --- MAIN GENERATION FUNCTION ---
def generate_resume(data: Dict[str, Any], template_path: str) -> Optional[io.BytesIO]:
    """Generates a resume by populating a .docx template with JSON data."""
    try:
        doc = Document(template_path)
    except Exception as e:
        st.error(f"Error opening template file '{template_path}': {e}")
        return None

    # --- Simple key-value replacements ---
    replacements = {
        "{NAME}": data.get("name", ""),
        "{CONTACT}": data.get("contact_number", ""),
        "{EMAIL}": data.get("email", ""),
        "{SKILLS}": ", ".join(data.get("skills", [])),
        "{LANGUAGES}": ", ".join(data.get("languages", [])),
        "{EDUCATION}": "\n\n".join(
            [
                f"{edu.get('degree', '')}\n{edu.get('institution', '')}\n{edu.get('year', '')}"
                for edu in data.get("education", [])
            ]
        ),
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_text_in_paragraph(p, key, value)

    # --- Dynamic Work Experience Table Population ---
    # This section is now corrected to avoid the AttributeError
    work_exp_placeholders = {"{COMPANYNAME}", "{JOBDESCRIPTION}", "{ACHIEVEMENTS}"}

    for table in doc.tables:
        # Find the single row that acts as a template
        template_info = None
        for i, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if any(ph in row_text for ph in work_exp_placeholders):
                template_info = {'row': row, 'index': i}
                break # Found our template, stop searching

        if not template_info:
            continue # This table is not for work experience, check the next one

        template_row = template_info['row']
        template_index = template_info['index']
        
        # Insert new rows for each work experience entry, starting from the bottom
        for experience in reversed(data.get("work_experience", [])):
            new_row_elem = deepcopy(template_row._element)
            # Insert the new row element right after the original template's position
            table._tbl.insert(template_index + 1, new_row_elem)
            new_row = table.rows[template_index + 1]

            # Populate the newly created row
            for cell in new_row.cells:
                for p in list(cell.paragraphs):
                    replace_text_in_paragraph(p, "{COMPANYNAME}", experience.get("company_name", ""))
                    replace_text_in_paragraph(p, "{DURATION}", experience.get("duration", ""))
                    replace_text_in_paragraph(p, "{JOBTITLE}", experience.get("job_title", ""))
                    
                    if "{JOBDESCRIPTION}" in p.text:
                        replace_with_bullet_points(p, "{JOBDESCRIPTION}", experience.get("job_description", []))
                    if "{ACHIEVEMENTS}" in p.text:
                        replace_with_bullet_points(p, "{ACHIEVEMENTS}", experience.get("achievements", []))
        
        # After adding all populated rows, remove the original template row
        table._tbl.remove(template_row._tr)
        break # Assume we only process one work experience table

    # Save to in-memory buffer for Streamlit
    try:
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
    except Exception as e:
        st.error(f"Error saving output file to memory: {e}")
        return None
